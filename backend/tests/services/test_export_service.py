from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

import fitz
from docx import Document

from app.schemas.exports import ExportRequest, SelectedStandard
from app.schemas.reports import ManualReviewHistoryEntry, ReportItem
from app.services.export_service import build_export_file


def _sample_card() -> ReportItem:
  return ReportItem(
    item_id="item-001",
    consistency_status="consistent",
    confidence_score=0.96,
    evidence="18.3 The Contractor shall finalise the EMP within 45 days.",
    reasoning="The evidence text directly states the deadline requirement.",
    document_references=["main_coc"],
    check_type="deadline",
    description="(PART 1) EMP finalisation timeline must be satisfied.",
    keywords=["EMP", "45 days"],
    source="pytest",
    severity="major",
    manual_verdict="needs_followup",
    manual_verdict_category="evidence_gap",
    manual_verdict_note="Need legal team confirmation for this clause.",
  )


def _sample_request(fmt: str) -> ExportRequest:
  return ExportRequest(
    report_id="rep_pytest_001",
    format=fmt,
    selected_standards=[
      SelectedStandard(
        standard_id="deadline",
        name="Deadline Compliance",
        priority=1,
      )
    ],
    card_ids=["item-001"],
  )


def _sample_manual_history() -> dict[str, list[ManualReviewHistoryEntry]]:
  return {
    "item-001": [
      ManualReviewHistoryEntry(
        history_id="mrh-002",
        report_id="rep_pytest_001",
        item_id="item-001",
        manual_verdict="needs_followup",
        manual_verdict_category="evidence_gap",
        manual_verdict_note="Need legal team confirmation for this clause.",
        edited_at=datetime(2026, 2, 26, 10, 30, tzinfo=timezone.utc),
      ),
      ManualReviewHistoryEntry(
        history_id="mrh-001",
        report_id="rep_pytest_001",
        item_id="item-001",
        manual_verdict="accepted",
        manual_verdict_category="false_positive",
        manual_verdict_note="Initial review accepted this item.",
        edited_at=datetime(2026, 2, 26, 9, 0, tzinfo=timezone.utc),
      ),
    ]
  }


def test_build_export_file_docx_contains_expected_content() -> None:
  request = _sample_request("docx")
  card = _sample_card()

  file_name, media_type, content = build_export_file(request, [card])

  assert file_name.endswith(".docx")
  assert media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
  assert content[:2] == b"PK"

  document = Document(BytesIO(content))
  paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
  table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

  assert "EPD Tender Analysis Report" in paragraph_text
  assert "Report ID: rep_pytest_001" in paragraph_text
  assert "Generated At (HKT):" in paragraph_text
  assert "Title: Consistency Review" in paragraph_text
  assert "Status: Consistent" in paragraph_text
  assert "Category: Deadline" in paragraph_text
  assert "Severity: Major" in paragraph_text
  assert "Confidence: 0.96" in paragraph_text
  assert "Referenced Sources: main_coc" in paragraph_text
  assert "Manual Review History:" in paragraph_text
  assert "No manual review history." in paragraph_text
  assert "Source: pytest" not in paragraph_text
  assert "Manual Verdict:" not in paragraph_text
  assert "Manual Note:" not in paragraph_text
  assert "(PART 1) EMP finalisation timeline must be satisfied." in paragraph_text
  assert "deadline" in table_text
  assert "Deadline Compliance" in table_text


def test_build_export_file_pdf_contains_expected_content() -> None:
  request = _sample_request("pdf")
  card = _sample_card()

  file_name, media_type, content = build_export_file(request, [card])

  assert file_name.endswith(".pdf")
  assert media_type == "application/pdf"
  assert content.startswith(b"%PDF")

  with fitz.open(stream=content, filetype="pdf") as document:
    extracted_text = "\n".join(page.get_text("text") for page in document)

  assert "EPD Tender Analysis Report" in extracted_text
  assert "rep_pytest_001" in extracted_text
  assert "Generated At (HKT):" in extracted_text
  assert "Title: Consistency Review" in extracted_text
  assert "Status: Consistent" in extracted_text
  assert "Category: Deadline" in extracted_text
  assert "Severity: Major" in extracted_text
  assert "Confidence: 0.96" in extracted_text
  assert "Referenced Sources: main_coc" in extracted_text
  assert "Manual Review History:" in extracted_text
  assert "No manual review history." in extracted_text
  assert "Source: pytest" not in extracted_text
  assert "Manual Verdict:" not in extracted_text
  assert "Manual Note:" not in extracted_text
  assert "EMP finalisation timeline" in extracted_text
  assert "Deadline Compliance" in extracted_text


def test_build_export_file_docx_includes_all_manual_review_history_entries() -> None:
  request = _sample_request("docx")
  card = _sample_card()
  manual_history = _sample_manual_history()

  _, _, content = build_export_file(request, [card], manual_history)
  document = Document(BytesIO(content))
  paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

  assert "Manual Review History:" in paragraph_text
  assert "1. Edited At (HKT):" in paragraph_text
  assert "2. Edited At (HKT):" in paragraph_text
  assert "Note: Need legal team confirmation for this clause." in paragraph_text
  assert "Note: Initial review accepted this item." in paragraph_text


def test_build_export_file_pdf_includes_all_manual_review_history_entries() -> None:
  request = _sample_request("pdf")
  card = _sample_card()
  manual_history = _sample_manual_history()

  _, _, content = build_export_file(request, [card], manual_history)
  with fitz.open(stream=content, filetype="pdf") as document:
    extracted_text = "\n".join(page.get_text("text") for page in document)
  normalized_text = " ".join(extracted_text.split())

  assert "Manual Review History:" in extracted_text
  assert "1. Edited At (HKT):" in extracted_text
  assert "2. Edited At (HKT):" in extracted_text
  assert "Need legal team confirmation for this clause." in normalized_text
  assert "Initial review accepted this item." in normalized_text


def test_build_export_file_docx_with_empty_card_ids_exports_no_cards() -> None:
  request = _sample_request("docx")
  request.card_ids = []
  card = _sample_card()

  _, _, content = build_export_file(request, [card])
  document = Document(BytesIO(content))
  paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

  assert "No cards selected for export." in paragraph_text
  assert "(PART 1) EMP finalisation timeline must be satisfied." not in paragraph_text


def test_build_export_file_pdf_with_empty_card_ids_exports_no_cards() -> None:
  request = _sample_request("pdf")
  request.card_ids = []
  card = _sample_card()

  _, _, content = build_export_file(request, [card])
  with fitz.open(stream=content, filetype="pdf") as document:
    extracted_text = "\n".join(page.get_text("text") for page in document)

  assert "No cards selected for export." in extracted_text
  assert "EMP finalisation timeline" not in extracted_text
