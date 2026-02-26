from __future__ import annotations

import re
from datetime import datetime, timedelta, timezone
from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.schemas.exports import ExportRequest
from app.schemas.reports import ReportItem


_FILENAME_RE = re.compile(r"[^a-zA-Z0-9._-]+")
_HONG_KONG_TZ = timezone(timedelta(hours=8), name="HKT")
_REVIEW_TITLE_BY_STATUS = {
  "consistent": "Consistency Review",
  "inconsistent": "Inconsistency Review",
  "unknown": "Consistency Review",
}


def _sanitize_file_token(value: str) -> str:
  normalized = _FILENAME_RE.sub("-", value).strip("-._")
  return normalized or "report"


def _ordered_cards(payload: ExportRequest, cards: list[ReportItem]) -> list[ReportItem]:
  if not cards or not payload.card_ids:
    return []

  by_id = {card.item_id: card for card in cards}
  ordered: list[ReportItem] = []
  seen: set[str] = set()

  for card_id in payload.card_ids:
    card = by_id.get(card_id)
    if card is None or card.item_id in seen:
      continue

    ordered.append(card)
    seen.add(card.item_id)

  return ordered


def _safe_text(text: str) -> str:
  return escape(text).replace("\n", "<br/>")


def _current_hong_kong_time() -> str:
  return datetime.now(_HONG_KONG_TZ).isoformat(timespec="seconds")


def _manual_review_value(value: str | None) -> str:
  return value if value else "N/A"


def _format_label(value: str) -> str:
  normalized = value.strip()
  if not normalized:
    return "N/A"

  parts = re.split(r"[\s_-]+", normalized)
  return " ".join(part.capitalize() for part in parts if part)


def _format_optional_label(value: str | None) -> str:
  if not value:
    return "N/A"
  return _format_label(value)


def _review_title(consistency_status: str) -> str:
  return _REVIEW_TITLE_BY_STATUS.get(consistency_status, "Consistency Review")


def _build_docx(payload: ExportRequest, cards: list[ReportItem]) -> bytes:
  ordered_cards = _ordered_cards(payload, cards)
  now = _current_hong_kong_time()

  document = Document()
  document.core_properties.title = f"EPD Tender Analysis - {payload.report_id}"

  document.add_heading("EPD Tender Analysis Report", level=0)
  document.add_paragraph(f"Report ID: {payload.report_id}")
  document.add_paragraph(f"Format: {payload.format.upper()}")
  document.add_paragraph(f"Generated At (HKT): {now}")

  document.add_heading("Selected Standards", level=1)
  if payload.selected_standards:
    standards = sorted(payload.selected_standards, key=lambda item: item.priority)
    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Priority"
    table.rows[0].cells[1].text = "Standard ID"
    table.rows[0].cells[2].text = "Name"

    for standard in standards:
      row = table.add_row().cells
      row[0].text = str(standard.priority)
      row[1].text = standard.standard_id
      row[2].text = standard.name
  else:
    document.add_paragraph("No selected standards.")

  document.add_heading("Cards", level=1)
  if not ordered_cards:
    document.add_paragraph("No cards selected for export.")

  for index, card in enumerate(ordered_cards, start=1):
    title = _review_title(card.consistency_status)
    document.add_heading(f"{index}. {title}", level=2)
    document.add_paragraph(f"Item ID: {card.item_id}")
    document.add_paragraph(f"Title: {title}")
    document.add_paragraph(f"Status: {_format_label(card.consistency_status)}")
    document.add_paragraph(f"Category: {_format_label(card.check_type)}")
    document.add_paragraph(f"Severity: {_format_label(card.severity)}")
    document.add_paragraph(f"Confidence: {card.confidence_score:.2f}")
    document.add_paragraph("Description:")
    document.add_paragraph(card.description)
    document.add_paragraph("Reasoning:")
    document.add_paragraph(card.reasoning)
    document.add_paragraph("Evidence:")
    document.add_paragraph(card.evidence)
    document.add_paragraph(f"Referenced Sources: {', '.join(card.document_references)}")
    if card.keywords:
      document.add_paragraph(f"Keywords: {', '.join(card.keywords)}")
    document.add_paragraph("Manual Review:")
    document.add_paragraph(
      f"Manual Verdict: {_format_optional_label(card.manual_verdict)} | "
      f"Manual Category: {_format_optional_label(card.manual_verdict_category)}"
    )
    document.add_paragraph(f"Manual Note: {_manual_review_value(card.manual_verdict_note)}")
    document.add_paragraph(f"Source: {card.source}")

  output = BytesIO()
  document.save(output)
  return output.getvalue()


def _build_pdf(payload: ExportRequest, cards: list[ReportItem]) -> bytes:
  ordered_cards = _ordered_cards(payload, cards)
  now = _current_hong_kong_time()

  output = BytesIO()
  doc = SimpleDocTemplate(
    output,
    pagesize=A4,
    leftMargin=40,
    rightMargin=40,
    topMargin=40,
    bottomMargin=40,
  )

  styles = getSampleStyleSheet()
  title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontSize=18, leading=22, spaceAfter=12)
  heading_style = ParagraphStyle("ReportHeading", parent=styles["Heading2"], fontSize=13, leading=16, spaceAfter=6)
  body_style = ParagraphStyle("ReportBody", parent=styles["BodyText"], fontSize=10, leading=14)
  meta_style = ParagraphStyle("ReportMeta", parent=styles["BodyText"], fontSize=9, leading=12, textColor=colors.HexColor("#334155"))

  story: list = []

  story.append(Paragraph("EPD Tender Analysis Report", title_style))
  story.append(Paragraph(f"Report ID: {_safe_text(payload.report_id)}", meta_style))
  story.append(Paragraph(f"Format: {_safe_text(payload.format.upper())}", meta_style))
  story.append(Paragraph(f"Generated At (HKT): {_safe_text(now)}", meta_style))
  story.append(Spacer(1, 12))

  story.append(Paragraph("Selected Standards", heading_style))
  if payload.selected_standards:
    rows = [["Priority", "Standard ID", "Name"]]
    for standard in sorted(payload.selected_standards, key=lambda item: item.priority):
      rows.append([str(standard.priority), standard.standard_id, standard.name])

    table = Table(rows, colWidths=[64, 160, 255])
    table.setStyle(
      TableStyle(
        [
          ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
          ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
          ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
          ("ALIGN", (0, 0), (0, -1), "CENTER"),
          ("VALIGN", (0, 0), (-1, -1), "TOP"),
          ("LEFTPADDING", (0, 0), (-1, -1), 6),
          ("RIGHTPADDING", (0, 0), (-1, -1), 6),
          ("TOPPADDING", (0, 0), (-1, -1), 4),
          ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]
      )
    )
    story.append(table)
  else:
    story.append(Paragraph("No selected standards.", body_style))

  story.append(Spacer(1, 12))
  story.append(Paragraph("Cards", heading_style))

  if not ordered_cards:
    story.append(Paragraph("No cards selected for export.", body_style))

  for index, card in enumerate(ordered_cards, start=1):
    title = _review_title(card.consistency_status)
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"{index}. {_safe_text(title)}", heading_style))
    story.append(Paragraph(f"Item ID: {_safe_text(card.item_id)}", meta_style))
    story.append(Paragraph(f"Title: {_safe_text(title)}", meta_style))
    story.append(Paragraph(f"Status: {_safe_text(_format_label(card.consistency_status))}", meta_style))
    story.append(Paragraph(f"Category: {_safe_text(_format_label(card.check_type))}", meta_style))
    story.append(Paragraph(f"Severity: {_safe_text(_format_label(card.severity))}", meta_style))
    story.append(Paragraph(f"Confidence: {card.confidence_score:.2f}", meta_style))
    story.append(Paragraph("Description:", meta_style))
    story.append(Paragraph(_safe_text(card.description), body_style))
    story.append(Paragraph("Reasoning:", meta_style))
    story.append(Paragraph(_safe_text(card.reasoning), body_style))
    story.append(Paragraph("Evidence:", meta_style))
    story.append(Paragraph(_safe_text(card.evidence), body_style))
    story.append(Paragraph(f"Referenced Sources: {_safe_text(', '.join(card.document_references))}", meta_style))
    if card.keywords:
      story.append(Paragraph(f"Keywords: {_safe_text(', '.join(card.keywords))}", meta_style))
    story.append(Paragraph("Manual Review:", meta_style))
    story.append(
      Paragraph(
        (
          f"Manual Verdict: {_safe_text(_format_optional_label(card.manual_verdict))} | "
          f"Manual Category: {_safe_text(_format_optional_label(card.manual_verdict_category))}"
        ),
        meta_style,
      )
    )
    story.append(Paragraph(f"Manual Note: {_safe_text(_manual_review_value(card.manual_verdict_note))}", body_style))
    story.append(Paragraph(f"Source: {_safe_text(card.source)}", meta_style))

  doc.build(story)
  return output.getvalue()


def build_export_file(payload: ExportRequest, cards: list[ReportItem]) -> tuple[str, str, bytes]:
  safe_report_id = _sanitize_file_token(payload.report_id)

  if payload.format == "pdf":
    media_type = "application/pdf"
    file_name = f"tender-analysis-{safe_report_id}.pdf"
    content = _build_pdf(payload, cards)
  else:
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_name = f"tender-analysis-{safe_report_id}.docx"
    content = _build_docx(payload, cards)

  return file_name, media_type, content
